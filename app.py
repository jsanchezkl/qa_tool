import streamlit as st
import pandas as pd
import asyncio
import re
import time
import io
from urllib.parse import urlparse, parse_qs
from typing import Dict, List, Any

# Playwright para la navegación asíncrona
from playwright.async_api import async_playwright, Request, Response, Page
from playwright.async_api import TimeoutError as PlaywrightTimeoutError

# Gemini para el análisis con IA
import google.generativeai as genai

# Docx para generar reportes en Word
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN Y CONSTANTES ---

# Configuración de la página de Streamlit
st.set_page_config(
    page_title="CSA | Auditoría Digital",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Argumentos para un lanzamiento más estable del navegador
BROWSER_LAUNCH_ARGS = ["--disable-http2", "--disable-quic", "--no-sandbox", "--disable-gpu"]
URL_DEFAULT = "https://www.nissan.com.co/"
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36"

# --- DEFINICIÓN DE FIRMAS DE ETIQUETAS ---

def get_tag_definitions() -> Dict[str, Dict[str, Any]]:
    """
    Centraliza todas las definiciones de etiquetas.
    Cada definición contiene:
    - pattern: Expresión regular para detectar la etiqueta en una URL.
    - category: El área a la que pertenece (Analytics, Ads, CMP, TMS, UX).
    - extractor: (Opcional) Una función para extraer el ID de cuenta/propiedad.
    """
    def generic_extractor(pattern: re.Pattern) -> callable:
        return lambda url: (m.group(1) if (m := pattern.search(url)) else "")

    return {
        # Tag Management Systems (TMS)
        "Google Tag Manager": {
            "pattern": re.compile(r"googletagmanager\.com/gtm\.js"),
            "category": "TMS",
            "extractor": generic_extractor(re.compile(r"[?&]id=(GTM-[A-Z0-9]+)"))
        },
        "Adobe Launch": {
            "pattern": re.compile(r"assets\.adobedtm\.com/"),
            "category": "TMS"
        },
        "Tealium iQ": {
            "pattern": re.compile(r"tags\.tiqcdn\.com/utag"),
            "category": "TMS"
        },
        # Consent Management Platforms (CMP)
        "Cookiebot": {
            "pattern": re.compile(r"consent\.cookiebot\.com/uc\.js"),
            "category": "CMP"
        },
        "OneTrust": {
            "pattern": re.compile(r"cdn\.cookielaw\.org/"),
            "category": "CMP"
        },
        "TrustArc": {
            "pattern": re.compile(r"consent\.trustarc\.com"),
            "category": "CMP"
        },
        # Analytics
        "Google Analytics 4 (Tag)": {
            "pattern": re.compile(r"gtag/js\?id=G-"),
            "category": "Analytics",
            "extractor": generic_extractor(re.compile(r"[?&]id=(G-[A-Z0-9]+)"))
        },
        "Google Analytics 4 (Event)": {
            "pattern": re.compile(r"google-analytics\.com/g/collect"),
            "category": "Analytics",
            "extractor": lambda url: (
                f"{parse_qs(urlparse(url).query).get('tid', ['N/A'])[0]} | "
                f"Evento: {parse_qs(urlparse(url).query).get('en', ['N/A'])[0]}"
            )
        },
        "Universal Analytics (Tag)": {
            "pattern": re.compile(r"google-analytics\.com/analytics\.js"),
            "category": "Analytics"
        },
        "Universal Analytics (Event)": {
            "pattern": re.compile(r"google-analytics\.com/collect\?v=1"),
            "category": "Analytics",
            "extractor": generic_extractor(re.compile(r"[?&]tid=(UA-[0-9\-]+)"))
        },
        "Adobe Analytics": {
            "pattern": re.compile(r"\.omtrdc\.net/b/ss/"),
            "category": "Analytics"
        },
        "Hotjar": {
            "pattern": re.compile(r"static\.hotjar\.com/c/hotjar-"),
            "category": "UX"
        },
        "Clarity": {
            "pattern": re.compile(r"clarity\.ms/tag/"),
            "category": "UX"
        },
        # Advertising & Retargeting
        "Meta Pixel (Facebook)": {
            "pattern": re.compile(r"connect\.facebook\.net/.*/fbevents\.js"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"[?&]id=([0-9]+)"))
        },
        "Meta Event (Facebook)": {
            "pattern": re.compile(r"facebook\.com/tr/\?"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"[?&]id=([0-9]+)"))
        },
        "TikTok Pixel": {
            "pattern": re.compile(r"analytics\.tiktok\.com/i18n/pixel"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"/pixel/([A-Z0-9]+)"))
        },
        "Google Ads Conversion": {
            "pattern": re.compile(r"google\.com/pagead/1p-user-list"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"/user-list/([A-Z0-9\-]+)/"))
        },
        "Google Ads Remarketing": {
            "pattern": re.compile(r"doubleclick\.net/pagead/viewthroughconversion"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"/viewthroughconversion/([0-9]+)/"))
        },
        "LinkedIn Insight": {
            "pattern": re.compile(r"snap\.licdn\.com/li\.lms-analytics/insight"),
            "category": "Ads"
        },
        "Bing UET": {
            "pattern": re.compile(r"bat\.bing\.com/bat\.js"),
            "category": "Ads",
            "extractor": generic_extractor(re.compile(r"[?&]ti=([0-9]+)"))
        },
        "Floodlight (DCM)": {
            "pattern": re.compile(r"doubleclick\.net/(fls|activity)"),
            "category": "Ads"
        },
        "Criteo": {
            "pattern": re.compile(r"static\.criteo\.net/js/ld/ld\.js"),
            "category": "Ads"
        },
    }

# --- LÓGICA DE SCRAPING Y ANÁLISIS ---

class DigitalAudit:
    """Encapsula la lógica de la auditoría."""
    def __init__(self, url: str, timeout_sec: int):
        self.url = self._format_url(url)
        self.timeout_ms = timeout_sec * 1000
        self.tag_definitions = get_tag_definitions()
        self.network_requests: List[Dict[str, Any]] = []
        self.found_tags: Dict[str, Dict[str, Any]] = {}

    def _format_url(self, url: str) -> str:
        """Asegura que la URL tenga un esquema."""
        if not url.startswith(('http://', 'https://')):
            return 'https://' + url
        return url

    async def run_audit(self):
        """Orquesta el proceso de auditoría con Playwright."""
        async with async_playwright() as p:
            try:
                browser = await p.chromium.launch(headless=True, args=BROWSER_LAUNCH_ARGS)
                context = await browser.new_context(user_agent=USER_AGENT)
                context.set_default_navigation_timeout(self.timeout_ms)
                context.on("request", self._handle_request)
                context.on("response", self._handle_response)
                page = await context.new_page()
                await self._navigate_and_interact(page)
                await browser.close()
            except PlaywrightTimeoutError:
                st.error(f"❌ Timeout ({self.timeout_ms / 1000}s) alcanzado. La auditoría puede estar incompleta.")
            except Exception as e:
                st.error(f"Ocurrió un error inesperado durante la auditoría: {e}")
                st.stop()

    async def _navigate_and_interact(self, page: Page):
        """Navega a la URL y espera a que la actividad de red se calme."""
        try:
            await page.goto(self.url, wait_until="networkidle", timeout=self.timeout_ms)
        except PlaywrightTimeoutError:
            st.warning("⚠️ 'networkidle' no se alcanzó. Se usará 'domcontentloaded'. Algunos tags podrían no ser detectados.")
            try:
                await page.goto(self.url, wait_until="domcontentloaded", timeout=self.timeout_ms)
            except PlaywrightTimeoutError:
                 st.error("❌ La página no pudo cargar ni siquiera el DOM básico en el tiempo asignado.")
                 return
        await asyncio.sleep(5)

    async def _handle_request(self, request: Request):
        """Almacena el tiempo de inicio de cada petición."""
        request.timing_info = {'start_time': time.perf_counter()}

    async def _handle_response(self, response: Response):
        """Analiza cada respuesta de red para identificar etiquetas."""
        request = response.request
        start_time = request.timing_info.get('start_time', time.perf_counter())
        load_time_ms = int((time.perf_counter() - start_time) * 1000)
        url = request.url
        for name, definition in self.tag_definitions.items():
            if definition["pattern"].search(url):
                account_id = ""
                if "extractor" in definition:
                    try:
                        account_id = definition["extractor"](url)
                    except Exception:
                        account_id = "Error al extraer"
                request_data = {
                    "Tag Name": name, "Category": definition["category"],
                    "Account ID / Details": account_id, "URL": url,
                    "Status": response.status, "Load (ms)": load_time_ms,
                    "Method": request.method,
                }
                self.network_requests.append(request_data)
                if name not in self.found_tags:
                    self.found_tags[name] = {"category": definition["category"], "count": 0}
                self.found_tags[name]["count"] += 1
                break

    def get_network_dataframe(self) -> pd.DataFrame:
        """Devuelve los datos de red como un DataFrame de Pandas."""
        if not self.network_requests:
            return pd.DataFrame()
        return pd.DataFrame(self.network_requests)

    def calculate_scores(self) -> Dict[str, Dict[str, Any]]:
        """Calcula las calificaciones y los criterios detallados para cada categoría."""
        scores = {
            "TMS": {"score": 1, "reason": "No se detectó un TMS.", "criteria": []},
            "Analytics": {"score": 1, "reason": "No se detectó GA4.", "criteria": []},
            "Ads": {"score": 1, "reason": "No se encontraron píxeles de publicidad.", "criteria": []},
            "CMP": {"score": 1, "reason": "No se detectó una plataforma de consentimiento.", "criteria": []},
            "Performance": {"score": 5, "reason": "El rendimiento de las etiquetas es excelente.", "criteria": []}
        }

        # TMS Score
        tms_found = any(t['category'] == 'TMS' for t in self.found_tags.values())
        scores["TMS"]["criteria"].append({"text": "Uso de un sistema de gestión de etiquetas (GTM, Adobe, etc.)", "pass": tms_found})
        if tms_found:
            scores["TMS"]["score"] = 5
            scores["TMS"]["reason"] = "Se detectó un TMS, una excelente práctica."
        
        # Analytics Score
        ga4_present = "Google Analytics 4 (Tag)" in self.found_tags or "Google Analytics 4 (Event)" in self.found_tags
        ua_present = "Universal Analytics (Tag)" in self.found_tags or "Universal Analytics (Event)" in self.found_tags
        scores["Analytics"]["criteria"].append({"text": "Implementación de Google Analytics 4 (GA4)", "pass": ga4_present})
        scores["Analytics"]["criteria"].append({"text": "Ausencia de Universal Analytics (obsoleto)", "pass": not ua_present})
        if ga4_present:
            scores["Analytics"]["score"] = 5
            scores["Analytics"]["reason"] = "¡Excelente! Se detectó Google Analytics 4."
            if ua_present:
                scores["Analytics"]["score"] = 4
                scores["Analytics"]["reason"] = "Se detectó GA4, pero también UA (obsoleto). Se recomienda migrar."
        elif ua_present:
            scores["Analytics"]["score"] = 2
            scores["Analytics"]["reason"] = "Alerta: Solo se detectó Universal Analytics (obsoleto)."

        # Ads Score
        ad_tags = [t for t, d in self.found_tags.items() if d['category'] == 'Ads']
        scores["Ads"]["criteria"].append({"text": "Presencia de al menos un píxel de publicidad", "pass": len(ad_tags) > 0})
        scores["Ads"]["criteria"].append({"text": "Cobertura en múltiples plataformas (>= 2 píxeles)", "pass": len(ad_tags) >= 2})
        if len(ad_tags) >= 4:
            scores["Ads"]["score"] = 5; scores["Ads"]["reason"] = "Excelente cobertura de píxeles en múltiples plataformas."
        elif len(ad_tags) >= 2:
            scores["Ads"]["score"] = 4; scores["Ads"]["reason"] = "Buena cobertura de píxeles de publicidad."
        elif len(ad_tags) > 0:
            scores["Ads"]["score"] = 3; scores["Ads"]["reason"] = "Se detectaron algunas etiquetas de publicidad."

        # CMP Score
        cmp_found = any(t['category'] == 'CMP' for t in self.found_tags.values())
        scores["CMP"]["criteria"].append({"text": "Uso de una Plataforma de Gestión de Consentimiento (CMP)", "pass": cmp_found})
        if cmp_found:
            scores["CMP"]["score"] = 5
            scores["CMP"]["reason"] = "Se detectó una CMP, excelente para la privacidad."

        # Performance Score
        df = self.get_network_dataframe()
        if not df.empty:
            avg_load = df["Load (ms)"].mean()
            errors = df[df["Status"] >= 400].shape[0]
            perf_score = 5
            reasons = []

            errors_pass = errors == 0
            scores["Performance"]["criteria"].append({"text": "Ausencia de errores en etiquetas (status 4xx/5xx)", "pass": errors_pass, "details": f"{errors} errores encontrados."})
            if not errors_pass: perf_score -= 2; reasons.append(f"{errors} etiquetas con errores.")

            avg_load_pass = avg_load <= 500
            scores["Performance"]["criteria"].append({"text": "Tiempo de carga promedio de etiquetas < 500ms", "pass": avg_load_pass, "details": f"Promedio: {avg_load:.0f} ms."})
            if not avg_load_pass: perf_score -= 1; reasons.append(f"Carga promedio elevada ({avg_load:.0f} ms).")
            
            if avg_load > 1000: perf_score -=1;
            
            scores["Performance"]["score"] = max(1, perf_score)
            if reasons: scores["Performance"]["reason"] = " ".join(reasons)
        
        return scores

# --- GENERACIÓN DE REPORTES ---

def generate_word_report(url: str, scores: Dict, summary_df: pd.DataFrame, gemini_report: str) -> io.BytesIO:
    """Crea un reporte de auditoría completo en un documento de Word."""
    document = Document()
    document.add_heading('Reporte Ejecutivo de Auditoría Digital', level=0)
    document.add_paragraph(f"URL Analizada: {url}")
    document.add_paragraph(f"Fecha de Auditoría: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    # --- Sección de Calificaciones ---
    document.add_heading('🏆 Dashboard de Calificaciones', level=1)
    avg_score = sum(s['score'] for s in scores.values()) / len(scores)
    document.add_paragraph(f"Calificación General Promedio: {avg_score:.1f} / 5")

    for cat, data in scores.items():
        document.add_heading(f"{cat}: {data['score']} / 5", level=2)
        p = document.add_paragraph()
        p.add_run('Resumen: ').bold = True
        p.add_run(data['reason'])
        
        p = document.add_paragraph()
        p.add_run('Criterios de Evaluación Detallados:').bold = True
        for criterion in data['criteria']:
            status_icon = "✅" if criterion['pass'] else "❌"
            document.add_paragraph(f"{status_icon} {criterion['text']}", style='List Bullet')
            if 'details' in criterion:
                 p_details = document.add_paragraph(f"   Detalle: {criterion['details']}", style='List Bullet 2')
                 p_details.paragraph_format.left_indent = Inches(0.5)

    # --- Sección de Resumen de Etiquetas ---
    document.add_heading('📄 Resumen de Etiquetas Encontradas', level=1)
    if not summary_df.empty:
        table = document.add_table(rows=1, cols=summary_df.shape[1], style='Table Grid')
        for i, col_name in enumerate(summary_df.columns):
            table.cell(0, i).text = col_name
        for index, row in summary_df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    else:
        document.add_paragraph("No se encontraron etiquetas conocidas.")

    # --- Sección de Análisis con IA ---
    document.add_heading('🧠 Análisis y Recomendaciones por IA (Gemini)', level=1)
    for line in gemini_report.split('\n'):
        if line.startswith('### '): document.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): document.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '): document.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('* '): document.add_paragraph(line.replace('* ', '', 1), style='List Bullet')
        else: document.add_paragraph(line)

    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f

def configure_gemini():
    """Configura la API de Gemini usando secretos de Streamlit."""
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        genai.configure(api_key=api_key)
        return True
    except (KeyError, AttributeError):
        st.error("🚨 Falta la API Key de Gemini. Configúrala en `secrets.toml`.")
        return False

def generate_gemini_report(scores: Dict, summary_df: pd.DataFrame) -> str:
    """Genera un reporte de consultoría usando el modelo Gemini."""
    if not configure_gemini():
        return "Análisis con IA no disponible por falta de API Key."

    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    prompt = f"""
    Actúa como un consultor experto en Analítica Digital y MarTech. He realizado una auditoría a un sitio web.
    Genera un reporte ejecutivo en español, claro, profesional y accionable.

    **CALIFICACIONES (1-5):**
    - **TMS:** {scores['TMS']['score']}/5. Razón: {scores['TMS']['reason']}
    - **Analytics:** {scores['Analytics']['score']}/5. Razón: {scores['Analytics']['reason']}
    - **Ads:** {scores['Ads']['score']}/5. Razón: {scores['Ads']['reason']}
    - **CMP:** {scores['CMP']['score']}/5. Razón: {scores['CMP']['reason']}
    - **Performance:** {scores['Performance']['score']}/5. Razón: {scores['Performance']['reason']}

    **RESUMEN DE ETIQUETAS:**
    {summary_df.to_string(index=False) if not summary_df.empty else "No se encontraron etiquetas."}

    **INSTRUCCIONES:**
    1.  **Título:** "Reporte Ejecutivo de Auditoría Digital".
    2.  **Resumen Ejecutivo:** Párrafo inicial resumiendo el estado general, la calificación promedio y los puntos clave.
    3.  **Análisis por Categoría:** Un apartado para cada una de las 5 categorías. Usa subtítulos (ej. "### 📊 Analítica Fundamental: Calificación X/5"). Explica la calificación y ofrece recomendaciones accionables.
    4.  **Conclusión y Próximos Pasos:** Finaliza con una conclusión y 2-3 próximos pasos prioritarios.
    5.  **Tono:** Profesional, didáctico y orientado a resultados. Usa formato Markdown.
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"❌ Error al contactar con la API de Gemini: {e}"

# --- INTERFAZ DE USUARIO (STREAMLIT) ---

def main():
    """Función principal que renderiza la aplicación Streamlit."""
    st.title("📊 CSA |Herramienta de Auditoría Digital")
    st.markdown("Analiza la implementación de etiquetas de analítica, publicidad y más en cualquier sitio web.")

    with st.form(key="audit_form"):
        url = st.text_input("Introduce la URL a auditar", value=URL_DEFAULT, placeholder="https://www.ejemplo.com")
        timeout_sec = st.slider("Timeout de navegación (segundos)", 15, 120, 45, 5)
        submit_button = st.form_submit_button(label="🚀 Auditar Sitio Web")

    if submit_button and url:
        with st.spinner(f"🕵️‍♂️ Auditando {url}... Esto puede tardar hasta {timeout_sec} segundos..."):
            audit = DigitalAudit(url, timeout_sec)
            asyncio.run(audit.run_audit())
            network_df = audit.get_network_dataframe()
            scores = audit.calculate_scores()
            found_tags_summary = pd.DataFrame([
                {"Etiqueta": name, "Categoría": data["category"], "Nº de Hits": data["count"]}
                for name, data in audit.found_tags.items()
            ]).sort_values(by="Categoría")
            gemini_report = generate_gemini_report(scores, found_tags_summary)

        st.success("✅ Auditoría completada.")
        
        # Botón de descarga del reporte en Word
        word_file_buffer = generate_word_report(url, scores, found_tags_summary, gemini_report)
        st.download_button(
            label="📄 Descargar Reporte Completo (.docx)",
            data=word_file_buffer,
            file_name=f"reporte_auditoria_{urlparse(url).netloc}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # --- Pestañas de Resultados ---
        tab1, tab2, tab3, tab4 = st.tabs(["🏆 Dashboard", "📄 Resumen Etiquetas", "📡 Registro de Red", "🧠 Análisis IA"])

        with tab1:
            st.header("🏆 Dashboard de Calificaciones")
            st.markdown("Evaluación automática de la madurez de la implementación digital del sitio (escala 1-5).")
            avg_score = sum(s['score'] for s in scores.values()) / len(scores)
            st.metric(label="Calificación General Promedio", value=f"{avg_score:.1f} / 5")
            st.progress(avg_score / 5)

            cols = st.columns(5)
            categories = list(scores.keys())
            icons = {"TMS": "🏷️", "Analytics": "📈", "Ads": "📢", "CMP": "🛡️", "Performance": "⚡"}
            
            for i, col in enumerate(cols):
                cat = categories[i]
                with col:
                    st.subheader(f"{icons[cat]} {cat}")
                    st.metric(label="Calificación", value=f"{scores[cat]['score']} / 5")
                    st.caption(scores[cat]['reason'])
                    with st.expander("Ver criterios de evaluación"):
                        for criterion in scores[cat]['criteria']:
                            status_icon = "✅" if criterion['pass'] else "❌"
                            st.markdown(f"{status_icon} **{criterion['text']}**")
                            if 'details' in criterion:
                                st.caption(criterion['details'])

        with tab2:
            st.header("📄 Resumen de Etiquetas Encontradas")
            if not found_tags_summary.empty: st.dataframe(found_tags_summary, use_container_width=True)
            else: st.warning("No se detectó ninguna etiqueta conocida.")

        with tab3:
            st.header("📡 Registro de Red Detallado")
            if not network_df.empty:
                st.dataframe(network_df, use_container_width=True, height=500)
                csv = network_df.to_csv(index=False).encode('utf-8')
                st.download_button("📥 Descargar CSV", csv, f"audit_{urlparse(url).netloc}.csv", "text/csv")
            else:
                st.warning("No se capturaron peticiones de red de etiquetas conocidas.")

        with tab4:
            st.header("🧠 Análisis y Recomendaciones por IA (Gemini)")
            with st.spinner("🤖 Gemini está preparando tu reporte de consultoría..."):
                st.markdown(gemini_report)

if __name__ == "__main__":
    main()
